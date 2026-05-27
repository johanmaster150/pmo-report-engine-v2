# ==========================================
# ENGINE LINUX — V5 Compilation (python-docx + LibreOffice)
# ==========================================
# Cross-platform replacement for Win32 COM automation.
# All document manipulation uses python-docx XML API.
# PDF/DOC conversions use LibreOffice headless.

import os
import re
import copy
import csv
import subprocess
import traceback

import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.table import Table
from docx.text.paragraph import Paragraph

# ============================================
# CONSTANTS
# ============================================
KNOWN_PROD_IDS = [
    '1.1', '1.2', '1.3', '1.4', '1.5', '1.6', '1.7',
    '2.1', '2.2', '2.3', '2.4',
    '3.1'
]

TOC_TITLES = [
    (1, "Informe Mensual"),
    (2, "Programa: Innovación como motor del desarrollo económico del Distrito"),
    (1, "Generalidades"),
    (2, "Objetivo General"),
    (2, "Proyectos por desarrollar"),
    (2, "Descripción informe"),
    (1, "Avance técnico Proyecto: Fortalecimiento del ecosistema CTi"),
    (2, "Avance Financiero y presupuestal"),
    (1, "Avance técnico Desarrollo de negocios de alto impacto"),
    (2, "Avance Financiero y presupuestal"),
    (1, "Avance técnico Proyecto: Centro para la Cuarta Revolución Industrial (C4RI)"),
    (1, "Avance técnico Proyecto: Centro para la Cuarta Revolución Industrial"),
    (2, "Avance Financiero y presupuestal"),
]

PRODUCT_KEYWORDS = [
    '#', 'n°', 'no.', 'num', 'producto', 'nombre', 'descripcion',
    'indicador', 'meta', 'actividad', 'avance', 'resultado',
    'ponderación', 'ponderacion', 'componente', 'hito',
    'entregable', 'evidencia', 'porcentaje', 'cumplimiento',
    'línea base', 'linea base'
]

PREFIX_PATTERN = re.compile(r'^\s*(\d+\.\d+)\b')


# ============================================
# UTILITIES
# ============================================

def convert_doc_to_docx(doc_path, output_dir=None):
    """Convert .doc to .docx using LibreOffice headless."""
    if output_dir is None:
        output_dir = os.path.dirname(doc_path)
    try:
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'docx', '--outdir', output_dir, doc_path],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode == 0:
            base = os.path.splitext(os.path.basename(doc_path))[0]
            converted = os.path.join(output_dir, base + '.docx')
            if os.path.exists(converted):
                print(f"  Converted {doc_path} -> {converted}")
                return converted
        print(f"  LibreOffice conversion failed: {result.stderr}")
    except FileNotFoundError:
        print("  WARNING: LibreOffice not found. Cannot convert .doc files.")
    except Exception as e:
        print(f"  Error converting .doc: {e}")
    return None


def convert_docx_to_format(docx_path, target_format='pdf', output_dir=None):
    """Convert DOCX to PDF or DOC using LibreOffice headless."""
    if output_dir is None:
        output_dir = os.path.dirname(docx_path)
    try:
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', target_format, '--outdir', output_dir, docx_path],
            capture_output=True, text=True, timeout=180
        )
        if result.returncode == 0:
            base = os.path.splitext(os.path.basename(docx_path))[0]
            converted = os.path.join(output_dir, base + '.' + target_format)
            if os.path.exists(converted):
                return converted
        print(f"  LibreOffice conversion to {target_format} failed: {result.stderr}")
    except FileNotFoundError:
        print(f"  WARNING: LibreOffice not found. Cannot convert to {target_format}.")
    except Exception as e:
        print(f"  Error converting to {target_format}: {e}")
    return None


def ensure_docx(file_path, temp_dir):
    """Ensure file is .docx. If .doc, convert it. Returns path to .docx file."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.doc':
        converted = convert_doc_to_docx(file_path, temp_dir)
        if converted:
            return converted
        raise ValueError(f"Could not convert .doc file: {file_path}")
    return file_path


def get_paragraph_text(paragraph):
    """Get clean text from a paragraph element."""
    return paragraph.text.strip() if paragraph.text else ""


def iter_block_items(doc):
    """
    Iterate over paragraphs and tables in document body order.
    Yields tuples of ('paragraph', Paragraph) or ('table', Table).
    """
    body = doc.element.body
    for child in body:
        if child.tag == qn('w:p'):
            yield ('paragraph', Paragraph(child, doc))
        elif child.tag == qn('w:tbl'):
            yield ('table', Table(child, doc))


def get_cell_text(table, row_idx, col_idx):
    """Safely get text from a table cell (0-indexed)."""
    try:
        return table.rows[row_idx].cells[col_idx].text.replace('\n', ' ').strip()
    except (IndexError, AttributeError):
        return ""


# ============================================
# PHASE 1: STRIP FORMATS
# ============================================

def strip_formats(doc, update_fn=None, task_id=None):
    """
    Phase 1: Remove all formatting artifacts from the merged document.
    - Remove highlights and shading from paragraphs (not in tables)
    - Reset Heading styles to Normal
    - Remove previous table numberings
    - Remove comments
    - Remove 'Tabla de Producto:' labels from V2
    """
    if update_fn:
        update_fn(task_id, 58, "V5 Fase 1: Eliminando formatos previos...")

    print("V5 FASE 1 (Linux): Stripping all formats...")

    # 1a. Remove comments from XML
    try:
        body = doc.element.body
        # Remove comment references and comments
        for tag_name in ['commentRangeStart', 'commentRangeEnd', 'commentReference']:
            for el in body.iter(qn(f'w:{tag_name}')):
                el.getparent().remove(el)
        # Remove comments part if it exists
        comments_parts = [
            rel for rel in doc.part.rels.values()
            if 'comments' in (rel.reltype or '')
        ]
        # We won't remove rels as that can corrupt, but comments are effectively orphaned
    except Exception as e:
        print(f"  V5 Fase 1 comment removal note: {e}")

    # 1b. Process paragraphs
    paragraphs_to_remove = []

    for para in doc.paragraphs:
        text = get_paragraph_text(para)

        # Check if paragraph is inside a table (skip table paragraphs for shading removal)
        is_in_table = para._element.getparent().tag == qn('w:tc')

        if not is_in_table:
            # Remove highlights from all runs
            for run in para.runs:
                run_elem = run._element
                rpr = run_elem.find(qn('w:rPr'))
                if rpr is not None:
                    # Remove highlight
                    highlight = rpr.find(qn('w:highlight'))
                    if highlight is not None:
                        rpr.remove(highlight)
                    # Remove shading from run
                    shd = rpr.find(qn('w:shd'))
                    if shd is not None:
                        rpr.remove(shd)

            # Remove paragraph-level shading
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                shd = pPr.find(qn('w:shd'))
                if shd is not None:
                    pPr.remove(shd)
                # Also check rPr inside pPr
                rPr = pPr.find(qn('w:rPr'))
                if rPr is not None:
                    shd = rPr.find(qn('w:shd'))
                    if shd is not None:
                        rPr.remove(shd)

        # Reset Heading styles to Normal
        style_name = ""
        try:
            style_name = para.style.name or ""
        except Exception:
            pass

        if 'Heading' in style_name or 'Título' in style_name or 'heading' in style_name.lower():
            try:
                para.style = doc.styles['Normal']
            except Exception:
                pass

        # Remove previous table numbering (e.g., "1.1.1 Tabla..." → "Tabla...")
        if text and re.match(r'^\s*\d+\.\d+\.\d+\s+[Tt]abla\b', text):
            new_text = re.sub(r'^\s*\d+\.\d+\.\d+\s+([Tt]abla)', r'\1', text)
            if new_text != text:
                _replace_paragraph_text(para, new_text)

        # Mark "Tabla de Producto:" paragraphs for removal
        if text and text.startswith("Tabla de Producto:"):
            paragraphs_to_remove.append(para)

    # Remove marked paragraphs
    for para in paragraphs_to_remove:
        _remove_paragraph(para)

    print("V5 Fase 1 (Linux): Format stripping complete.")


def _replace_paragraph_text(para, new_text):
    """Replace all text in a paragraph while keeping first run's formatting."""
    if para.runs:
        # Keep first run's formatting, clear all runs
        first_run_props = copy.deepcopy(para.runs[0]._element.find(qn('w:rPr')))
        for run in para.runs:
            run._element.getparent().remove(run._element)
        # Add new run with original formatting
        new_run = para.add_run(new_text)
        if first_run_props is not None:
            new_run._element.insert(0, first_run_props)
    else:
        para.text = new_text


def _remove_paragraph(para):
    """Remove a paragraph element from the document."""
    p = para._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


# ============================================
# PHASE 2: REBUILD TITLES
# ============================================

def rebuild_titles(doc, update_fn=None, task_id=None):
    """
    Phase 2: Detect products in tables and rebuild section titles.
    Returns list of product_sections: [{prod_id, name, table_index, body_index}]
    """
    if update_fn:
        update_fn(task_id, 65, "V5 Fase 2: Detectando y reconstruyendo títulos de productos...")

    print("V5 FASE 2 (Linux): Rebuilding titles/subtitles...")

    product_sections = []
    existing_ids = set()
    body = doc.element.body

    # --- Build index of all tables with their body position ---
    tables_info = []  # [(table_obj, body_index)]
    for idx, child in enumerate(body):
        if child.tag == qn('w:tbl'):
            tbl_obj = Table(child, doc)
            tables_info.append((tbl_obj, idx))
            # Apply borders and autofit
            _apply_table_formatting(tbl_obj)

    # --- Strategy A: Header-keyword detection ---
    for tbl_obj, body_idx in tables_info:
        try:
            num_rows = len(tbl_obj.rows)
            num_cols = len(tbl_obj.columns)
            if num_cols < 2:
                continue

            # Check headers
            h1 = get_cell_text(tbl_obj, 0, 0).lower()
            h2 = get_cell_text(tbl_obj, 0, 1).lower()
            h3 = get_cell_text(tbl_obj, 0, 2).lower() if num_cols >= 3 else ''
            all_headers = f"{h1} {h2} {h3}"

            if not any(kw in all_headers for kw in PRODUCT_KEYWORDS):
                continue

            # Scan rows for prod_id
            prod_id = None
            prod_name = None
            for row_idx in range(1, min(num_rows, 10)):
                val1 = get_cell_text(tbl_obj, row_idx, 0)
                val2 = get_cell_text(tbl_obj, row_idx, 1) if num_cols >= 2 else ''

                m1 = PREFIX_PATTERN.match(val1)
                m2 = PREFIX_PATTERN.match(val2)

                if m1 and val2 and not PREFIX_PATTERN.match(val2):
                    prod_id = m1.group(1)
                    prod_name = val2.strip()
                    break
                elif m2 and val1 and not PREFIX_PATTERN.match(val1):
                    prod_id = m2.group(1)
                    prod_name = val1.strip()
                    break
                elif m1:
                    prod_id = m1.group(1)
                    prod_name = val2 if val2 else val1
                    break

            if prod_id and prod_name and len(prod_name) > 2 and prod_id not in existing_ids:
                existing_ids.add(prod_id)
                product_sections.append({
                    "prod_id": prod_id,
                    "name": prod_name[:150],
                    "body_index": body_idx,
                    "table_element": tbl_obj._tbl
                })

        except Exception as e:
            pass

    print(f"V5 Fase 2 Strategy A (Linux): Found {len(product_sections)} products via header keywords.")

    # --- Strategy B: Search for known IDs in paragraphs ---
    # Refresh body index after potential changes
    body_elements = list(body)

    for known_id in KNOWN_PROD_IDS:
        if known_id in existing_ids:
            continue

        # Search paragraphs for this ID
        found_idx = None
        for idx, child in enumerate(body_elements):
            if child.tag == qn('w:p'):
                p_text = Paragraph(child, doc).text.strip()
                if known_id in p_text:
                    found_idx = idx
                    break

        if found_idx is not None:
            # Find nearest table after this position
            nearest_table = None
            nearest_body_idx = None
            for tbl_obj, tbl_body_idx in tables_info:
                if tbl_body_idx > found_idx:
                    # Check not already claimed
                    already_claimed = any(
                        sec['body_index'] == tbl_body_idx
                        for sec in product_sections
                    )
                    if not already_claimed:
                        nearest_table = tbl_obj
                        nearest_body_idx = tbl_body_idx
                        break

            if nearest_table and nearest_body_idx:
                # Try to get product name from table
                prod_name_b = _extract_product_name_from_table(nearest_table)
                if not prod_name_b:
                    prod_name_b = f"Producto {known_id}"

                product_sections.append({
                    "prod_id": known_id,
                    "name": prod_name_b,
                    "body_index": nearest_body_idx,
                    "table_element": nearest_table._tbl
                })
                existing_ids.add(known_id)
                print(f"  V5 Strategy B (Linux): Found '{known_id}'")

    print(f"V5 Fase 2 Total (Linux): {len(product_sections)} products found.")

    # --- Sort by body position ---
    product_sections.sort(key=lambda x: x['body_index'])

    # --- Insert Heading 2 titles before each product table ---
    offset = 0  # Track insertions shifting indices
    for ps in product_sections:
        try:
            tbl_elem = ps['table_element']
            title_text = f"{ps['prod_id']}. {ps['name']}"

            # Create a new paragraph element for the heading
            heading_p = _create_heading_paragraph(doc, title_text, level=2)

            # Insert before the table in the body
            body.insert(list(body).index(tbl_elem), heading_p)
            offset += 1

        except Exception as e:
            print(f"  V5 Error inserting heading for {ps['prod_id']}: {e}")

    # --- Apply Heading 1 to main section titles ---
    main_headings = [
        "Avance técnico Proyecto: Fortalecimiento del ecosistema CTi",
        "Avance técnico Desarrollo de negocios de alto impacto",
        "Avance técnico Proyecto: Centro para la Cuarta Revolución Industrial",
        "Generalidades",
        "Informe Mensual",
    ]

    for heading_text in main_headings:
        for para in doc.paragraphs:
            if heading_text.lower() in para.text.lower():
                try:
                    para.style = doc.styles['Heading 1']
                except KeyError:
                    try:
                        para.style = doc.styles['Título 1']
                    except KeyError:
                        # Manual formatting fallback
                        for run in para.runs:
                            run.bold = True
                            run.font.size = Pt(16)
                break

    print("V5 Fase 2 (Linux): Title reconstruction complete.")
    return product_sections


def _extract_product_name_from_table(table):
    """Try to extract a descriptive product name from a table."""
    num_rows = len(table.rows)
    num_cols = len(table.columns)
    for row_idx in range(min(num_rows, 10)):
        for col_idx in range(min(num_cols, 4)):
            text = get_cell_text(table, row_idx, col_idx)
            if text and len(text) > 5 and not PREFIX_PATTERN.match(text):
                low = text.lower()
                if not any(kw in low for kw in ['#', 'n°', 'indicador', 'meta', 'actividad']):
                    return text[:100]
    return None


def _apply_table_formatting(table):
    """Apply borders and alignment to a table using XML."""
    try:
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
            tbl.insert(0, tblPr)

        # Set borders
        borders_xml = (
            f'<w:tblBorders {nsdecls("w")}>'
            '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '</w:tblBorders>'
        )
        # Remove existing borders
        existing = tblPr.find(qn('w:tblBorders'))
        if existing is not None:
            tblPr.remove(existing)
        tblPr.append(parse_xml(borders_xml))

        # Set table width to 100%
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
            tblPr.append(tblW)
        else:
            tblW.set(qn('w:w'), '5000')
            tblW.set(qn('w:type'), 'pct')

        # Center alignment
        jc = tblPr.find(qn('w:jc'))
        if jc is None:
            jc = parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>')
            tblPr.append(jc)
        else:
            jc.set(qn('w:val'), 'center')

    except Exception as e:
        pass


def _create_heading_paragraph(doc, text, level=2):
    """Create a heading paragraph XML element."""
    p = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr>'
        f'    <w:pStyle w:val="Heading{level}"/>'
        f'  </w:pPr>'
        f'  <w:r>'
        f'    <w:rPr>'
        f'      <w:b/>'
        f'      <w:sz w:val="{28 if level == 1 else 24}"/>'
        f'      <w:color w:val="1F3864"/>'
        f'    </w:rPr>'
        f'    <w:t>{text}</w:t>'
        f'  </w:r>'
        f'</w:p>'
    )
    return p


# ============================================
# PHASE 3: TABLE NUMBERING
# ============================================

def number_tables(doc, product_sections, update_fn=None, task_id=None):
    """
    Phase 3: Number tables per product section as X.Y.Z format.
    """
    if update_fn:
        update_fn(task_id, 75, "V5 Fase 3: Numerando tablas por sección (X.Y.Z)...")

    print("V5 FASE 3 (Linux): Table numbering...")

    body = doc.element.body
    body_elements = list(body)

    # Build section ranges based on body indices
    # First, find current positions of product table elements
    section_positions = []
    for ps in product_sections:
        tbl_elem = ps['table_element']
        try:
            idx = body_elements.index(tbl_elem)
            section_positions.append((ps['prod_id'], idx))
        except ValueError:
            # Table might have shifted due to insertions; search for it
            for i, child in enumerate(body_elements):
                if child is tbl_elem:
                    section_positions.append((ps['prod_id'], i))
                    break

    section_positions.sort(key=lambda x: x[1])

    # Build ranges
    section_ranges = []
    for i, (pid, start_idx) in enumerate(section_positions):
        end_idx = section_positions[i + 1][1] if i + 1 < len(section_positions) else len(body_elements)
        section_ranges.append((pid, start_idx, end_idx))

    def get_prod_id_for_index(idx):
        for pid, s, e in section_ranges:
            if s <= idx < e:
                return pid
        return None

    # Iterate all tables and assign numbering
    table_seq_by_prod = {}
    tables_numbered = 0

    for child_idx, child in enumerate(body_elements):
        if child.tag != qn('w:tbl'):
            continue

        tbl_obj = Table(child, doc)
        current_prod = get_prod_id_for_index(child_idx)
        if not current_prod:
            continue

        seq = table_seq_by_prod.get(current_prod, 0) + 1
        table_seq_by_prod[current_prod] = seq

        table_label = f"Tabla {current_prod}.{seq}"

        # Get descriptive name from header
        desc = ""
        try:
            num_cols = len(tbl_obj.columns)
            if num_cols >= 2:
                h1 = get_cell_text(tbl_obj, 0, 0)
                h2 = get_cell_text(tbl_obj, 0, 1)
                if len(h2) > len(h1) and len(h2) > 3:
                    desc = h2[:80]
                elif len(h1) > 3:
                    desc = h1[:80]
        except Exception:
            pass

        full_label = f"{table_label}: {desc}" if desc else table_label

        # Insert label paragraph before the table
        try:
            label_p = _create_table_label_paragraph(full_label)
            body.insert(child_idx + tables_numbered, label_p)
            tables_numbered += 1
        except Exception as e:
            print(f"  V5 Error labeling table: {e}")

    print(f"V5 Fase 3 (Linux): Numbered {tables_numbered} tables across {len(table_seq_by_prod)} products.")

    # Also renumber standalone "Tabla N" captions
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and re.match(r'^\s*[Tt]abla\s+\d+', text) and not re.match(r'^\s*[Tt]abla\s+\d+\.\d+', text):
            # Find which section this paragraph belongs to
            p_elem = para._element
            try:
                p_idx = list(body).index(p_elem)
            except ValueError:
                continue
            current_prod = get_prod_id_for_index(p_idx)
            if current_prod:
                seq = table_seq_by_prod.get(current_prod, 0) + 1
                table_seq_by_prod[current_prod] = seq
                new_text = re.sub(
                    r'^\s*[Tt]abla\s+\d+',
                    f'Tabla {current_prod}.{seq}',
                    text
                )
                if new_text != text:
                    _replace_paragraph_text(para, new_text)

    return table_seq_by_prod


def _create_table_label_paragraph(text):
    """Create an italic, centered label paragraph for tables."""
    p = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr>'
        f'    <w:jc w:val="center"/>'
        f'  </w:pPr>'
        f'  <w:r>'
        f'    <w:rPr>'
        f'      <w:i/>'
        f'      <w:sz w:val="18"/>'
        f'      <w:color w:val="404040"/>'
        f'    </w:rPr>'
        f'    <w:t>{text}</w:t>'
        f'  </w:r>'
        f'</w:p>'
    )
    return p


# ============================================
# PHASE 4: EMPTY SPACE OPTIMIZATION
# ============================================

def optimize_empty_space(doc, update_fn=None, task_id=None):
    """Remove consecutive empty paragraphs (keep max 2)."""
    if update_fn:
        update_fn(task_id, 80, "V5: Optimizando espacios vacíos...")

    body = doc.element.body
    to_remove = []
    consecutive_empty = 0

    for child in body:
        if child.tag == qn('w:p'):
            text = Paragraph(child, doc).text.strip()
            if not text:
                consecutive_empty += 1
                if consecutive_empty > 2:
                    to_remove.append(child)
            else:
                consecutive_empty = 0
        else:
            consecutive_empty = 0

    for elem in to_remove:
        body.remove(elem)

    print(f"V5 (Linux): Removed {len(to_remove)} excess empty paragraphs.")


# ============================================
# PHASE 5: CUSTOM TOC
# ============================================

def build_custom_toc(doc, product_sections, update_fn=None, task_id=None):
    """Build a custom Table of Contents based on detected headings and products."""
    if update_fn:
        update_fn(task_id, 83, "V5: Construyendo tabla de contenido...")

    print("V5 (Linux): Building custom TOC...")

    toc_entries = []
    avance_financiero_count = 0

    # Search for known TOC titles in paragraphs
    for level, title_search in TOC_TITLES:
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if title_search.lower() in text.lower():
                if "Avance Financiero y presupuestal" in title_search:
                    avance_financiero_count += 1
                    if avance_financiero_count > 1:
                        # Need to find the Nth occurrence
                        count = 0
                        found = False
                        for p2_idx, p2 in enumerate(doc.paragraphs):
                            if title_search.lower() in p2.text.lower():
                                count += 1
                                if count == avance_financiero_count:
                                    toc_entries.append({
                                        "level": level,
                                        "title": title_search,
                                        "page": "—",
                                        "order": p2_idx
                                    })
                                    found = True
                                    break
                        if found:
                            break
                        continue

                toc_entries.append({
                    "level": level,
                    "title": title_search,
                    "page": "—",
                    "order": para_idx
                })
                break

    # Add products to TOC
    for ps in product_sections:
        toc_entries.append({
            "level": 2,
            "title": f"Producto {ps['prod_id']}: {ps['name']}",
            "page": "—",
            "order": ps.get('body_index', 9999)
        })

    toc_entries.sort(key=lambda x: x["order"])

    if not toc_entries:
        print("  No TOC entries found.")
        return

    # Build TOC text
    toc_lines = []
    for entry in toc_entries:
        indent = "    " if entry["level"] == 2 else ""
        title_display = entry["title"]
        page_str = str(entry["page"])
        dots = '.' * max(3, 70 - len(indent) - len(title_display) - len(page_str))
        toc_lines.append((entry["level"], f"{indent}{title_display} {dots} {page_str}"))

    # Find placeholder "Tabla de Contenido" in the document
    body = doc.element.body
    toc_placeholder_para = None
    for para in doc.paragraphs:
        if re.search(r'tabla\s+de\s+contenido', para.text, re.IGNORECASE):
            toc_placeholder_para = para
            break

    if toc_placeholder_para:
        # Format the placeholder heading
        toc_placeholder_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in toc_placeholder_para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(16)
            run.bold = True

        # Insert TOC entries after the placeholder
        insert_after = toc_placeholder_para._element
        for level, line in reversed(toc_lines):
            toc_p = _create_toc_entry_paragraph(line, level)
            insert_after.addnext(toc_p)

        # Add blank paragraph after heading
        blank_p = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
        toc_placeholder_para._element.addnext(blank_p)

        print(f"V5 TOC (Linux): Inserted at placeholder with {len(toc_entries)} entries.")
    else:
        # Prepend at document start
        # Insert in reverse order at document start
        for level, line in reversed(toc_lines):
            toc_p = _create_toc_entry_paragraph(line, level)
            body.insert(0, toc_p)

        # Insert blank line
        blank_p = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
        body.insert(0, blank_p)

        # Insert title
        title_p = parse_xml(
            f'<w:p {nsdecls("w")}>'
            f'  <w:pPr><w:jc w:val="center"/></w:pPr>'
            f'  <w:r>'
            f'    <w:rPr><w:b/><w:sz w:val="32"/><w:rFonts w:ascii="Arial" w:hAnsi="Arial"/></w:rPr>'
            f'    <w:t>Tabla de Contenido</w:t>'
            f'  </w:r>'
            f'</w:p>'
        )
        body.insert(0, title_p)

        print(f"V5 TOC (Linux): Prepended at start with {len(toc_entries)} entries.")


def _create_toc_entry_paragraph(text, level):
    """Create a formatted TOC entry paragraph."""
    font_size = 22 if level == 1 else 20  # In half-points
    bold_tag = '<w:b/>' if level == 1 else ''
    p = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:r>'
        f'    <w:rPr>'
        f'      {bold_tag}'
        f'      <w:sz w:val="{font_size}"/>'
        f'      <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>'
        f'    </w:rPr>'
        f'    <w:t xml:space="preserve">{text}</w:t>'
        f'  </w:r>'
        f'</w:p>'
    )
    return p


# ============================================
# PHASE 6: UNIFIED TYPOGRAPHY
# ============================================

def apply_typography(doc, font_family, update_fn=None, task_id=None):
    """Apply a unified font family to the entire document."""
    if font_family == "default":
        return

    if update_fn:
        update_fn(task_id, 88, f"V5: Aplicando tipografía {font_family}...")

    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = font_family

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = font_family

    print(f"V5 (Linux): Typography set to {font_family}.")


# ============================================
# PHASE 7: LINK EXTRACTION
# ============================================

def extract_links(doc, output_dir, product_sections, check_cti_fn, update_fn=None, task_id=None):
    """Extract hyperlinks from document and build evidence report."""
    if update_fn:
        update_fn(task_id, 89, "V5: Extrayendo enlaces y evidencias...")

    import requests as req_lib

    evidencias_dir = os.path.join(output_dir, "Evidencias_Descargadas")
    os.makedirs(evidencias_dir, exist_ok=True)

    links_data = []
    link_counters = {}

    # Extract hyperlinks from document XML
    body = doc.element.body
    all_hyperlinks = body.findall('.//' + qn('w:hyperlink'))

    rels = doc.part.rels

    for hl in all_hyperlinks:
        r_id = hl.get(qn('r:id'))
        if not r_id or r_id not in rels:
            continue

        rel = rels[r_id]
        url = rel.target_ref if hasattr(rel, 'target_ref') else str(rel._target)

        if not url or not url.startswith('http'):
            continue

        # Get display text
        text_parts = []
        for t_el in hl.findall('.//' + qn('w:t')):
            if t_el.text:
                text_parts.append(t_el.text)
        text_to_display = ' '.join(text_parts) or url

        # Determine product section
        # Simple heuristic: check if any product name appears in nearby text
        current_prod = None
        # Try to match by position in document
        parent = hl.getparent()
        if parent is not None:
            try:
                parent_idx = list(body).index(parent)
                for ps in reversed(product_sections):
                    if parent_idx >= ps.get('body_index', 0):
                        current_prod = ps['prod_id']
                        break
            except ValueError:
                pass

        if current_prod:
            link_counters[current_prod] = link_counters.get(current_prod, 0) + 1
            seq_id = f"{current_prod}.{link_counters[current_prod]}"
        else:
            link_counters["general"] = link_counters.get("general", 0) + 1
            seq_id = f"G.{link_counters['general']}"

        cti_status = check_cti_fn(url)

        status = "Fallido (Requiere Autenticación)"
        try:
            res = req_lib.get(url, timeout=5, stream=True)
            if res.status_code == 200 and 'text/html' not in res.headers.get('Content-Type', ''):
                from werkzeug.utils import secure_filename
                filename = secure_filename(url.split('/')[-1])
                if not filename:
                    filename = f"evidencia_{len(links_data)}.file"
                filepath = os.path.join(evidencias_dir, filename)
                with open(filepath, 'wb') as f:
                    for chunk in res.iter_content(chunk_size=8192):
                        if chunk:
                            f.write(chunk)
                status = "Descargado"
        except Exception:
            pass

        links_data.append([seq_id, url, text_to_display, cti_status, status])

    csv_filename = None
    docx_filename = None

    if links_data:
        csv_filename = "00_Reporte_Enlaces_Evidencias_V5.csv"
        docx_filename = "02_Listado_Enlaces_y_Evidencias_V5.docx"
        csv_path = os.path.join(output_dir, csv_filename)
        docx_path = os.path.join(output_dir, docx_filename)

        # Write CSV
        with open(csv_path, 'w', newline='', encoding='utf-8-sig') as f:
            writer = csv.writer(f)
            writer.writerow(["#", "Enlace (URL)", "Texto en Documento", "¿En ruta CTI?", "Estado de Descarga"])
            writer.writerows(links_data)

        # Build links DOCX
        links_doc = Document()
        links_doc.add_heading('Listado de Enlaces y Evidencias (V5 Linux)', level=1)

        table = links_doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        headers = ["#", "Enlace (URL)", "Texto en Documento", "¿En ruta CTI?", "Estado"]
        for col_idx, header_text in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            cell.text = header_text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            shading_elm = cell._element.get_or_add_tcPr()
            shading = shading_elm.makeelement(qn('w:shd'), {
                qn('w:fill'): '2F5496',
                qn('w:val'): 'clear'
            })
            shading_elm.append(shading)

        for row_data in links_data:
            row_cells = table.add_row().cells
            for col_idx, value in enumerate(row_data):
                row_cells[col_idx].text = str(value)
                for paragraph in row_cells[col_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        links_doc.save(docx_path)
        print(f"V5 Links (Linux): Extracted {len(links_data)} links.")

    return csv_filename, docx_filename


# ============================================
# MERGE ENGINE - Insert files into master
# ============================================

def merge_documents(master_path, project_files, temp_dir, update_fn=None, task_id=None):
    """
    Merge multiple DOCX files into the master template.
    Replaces Win32's InsertFile with python-docx body element copying.
    Returns the merged Document object.
    """
    if update_fn:
        update_fn(task_id, 47, "V5 Fase 0: Preparando plantilla maestra...")

    print("V5 (Linux): Opening master template...")
    master_doc = Document(master_path)

    # Clear master template body content (keep styles, headers, footers via section properties)
    body = master_doc.element.body

    # Preserve section properties (last section = headers/footers/page setup)
    sect_props = body.findall(qn('w:sectPr'))
    last_sect_pr = None
    if sect_props:
        last_sect_pr = copy.deepcopy(sect_props[-1])

    # Clear all body content
    for child in list(body):
        body.remove(child)

    # Restore section properties
    if last_sect_pr is not None:
        body.append(last_sect_pr)

    print("V5 (Linux): Master template content cleared.")

    # Insert project files
    total_p = len(project_files)
    for i, doc_path in enumerate(project_files):
        pct = 48 + int(10 * (i / total_p))
        if update_fn:
            update_fn(task_id, pct, f"V5 Fase 0: Insertando {os.path.basename(doc_path)}...")

        print(f"V5 Inserting (Linux): {os.path.basename(doc_path)}")

        try:
            # Ensure .docx format
            docx_path = ensure_docx(doc_path, temp_dir)
            src_doc = Document(docx_path)

            # Insert section break before each file (except first)
            if i > 0:
                # Add a section break paragraph
                break_p = parse_xml(
                    f'<w:p {nsdecls("w")}>'
                    f'  <w:pPr>'
                    f'    <w:sectPr>'
                    f'      <w:type w:val="nextPage"/>'
                    f'    </w:sectPr>'
                    f'  </w:pPr>'
                    f'</w:p>'
                )
                if last_sect_pr is not None:
                    body.insert(list(body).index(last_sect_pr), break_p)
                else:
                    body.append(break_p)

            # Copy all body elements from source document
            for child in src_doc.element.body:
                # Skip section properties from source (we keep master's)
                if child.tag == qn('w:sectPr'):
                    continue
                # Deep copy the element
                new_el = copy.deepcopy(child)
                # Insert before the last section properties
                if last_sect_pr is not None and last_sect_pr in body:
                    body.insert(list(body).index(last_sect_pr), new_el)
                else:
                    body.append(new_el)

            # Handle relationships (images, etc.) from source document
            _copy_relationships(src_doc, master_doc)

        except Exception as e:
            print(f"V5 Error inserting {doc_path}: {e}")
            traceback.print_exc()

    print(f"V5 (Linux): All {total_p} files inserted.")
    return master_doc


def _copy_relationships(src_doc, dest_doc):
    """Copy image and other relationships from source to destination document."""
    try:
        for rel_id, rel in src_doc.part.rels.items():
            if rel.is_external:
                continue
            # Copy image parts
            if 'image' in (rel.reltype or ''):
                try:
                    image_part = rel.target_part
                    # Add the image to dest
                    new_rel_id = dest_doc.part.relate_to(image_part, rel.reltype)
                    # Update references in the body (done by deep copy, so element refs should work)
                except Exception:
                    pass
    except Exception:
        pass


# ============================================
# MAIN COMPILATION PIPELINE
# ============================================

def compile_unified_v5_linux(master_path, project_files, output_path,
                              font_family="default", export_format="word",
                              task_id=None, update_fn=None, check_cti_fn=None,
                              output_folder=None):
    """
    Main V5 compilation pipeline for Linux.
    Equivalent to compile_unified_word_v5() but without any Win32 COM.
    """
    if update_fn:
        update_fn(task_id, 45, "V5 Motor Linux: Iniciando pipeline...")

    temp_dir = os.path.dirname(output_path)

    try:
        # MERGE: Insert all project files into master template
        doc = merge_documents(master_path, project_files, temp_dir, update_fn, task_id)

        # PHASE 1: Strip formats
        strip_formats(doc, update_fn, task_id)

        # PHASE 2: Rebuild titles
        product_sections = rebuild_titles(doc, update_fn, task_id)

        # PHASE 3: Number tables
        number_tables(doc, product_sections, update_fn, task_id)

        # PHASE 4: Empty space optimization
        optimize_empty_space(doc, update_fn, task_id)

        # PHASE 5: Custom TOC
        build_custom_toc(doc, product_sections, update_fn, task_id)

        # PHASE 6: Typography
        apply_typography(doc, font_family, update_fn, task_id)

        # PHASE 7: Link extraction
        if update_fn:
            update_fn(task_id, 89, "V5: Extrayendo enlaces y evidencias...")
        csv_filename, links_docx_filename = None, None
        if check_cti_fn and output_folder:
            csv_filename, links_docx_filename = extract_links(
                doc, output_folder, product_sections, check_cti_fn, update_fn, task_id
            )

        # PHASE 8: Save
        if update_fn:
            update_fn(task_id, 95, "V5: Guardando documento ensamblado...")

        files_saved = {}

        # Always save as DOCX first (needed for conversions)
        docx_output = output_path
        if not docx_output.endswith('.docx'):
            docx_output = output_path.rsplit('.', 1)[0] + '.docx'

        doc.save(docx_output)
        print(f"V5 (Linux): DOCX saved to {docx_output}")

        if export_format in ["word", "both"]:
            files_saved["docx"] = os.path.basename(docx_output)

        if export_format in ["pdf", "both"]:
            if update_fn:
                update_fn(task_id, 97, "V5: Convirtiendo a PDF via LibreOffice...")
            pdf_path = convert_docx_to_format(docx_output, 'pdf', output_folder)
            if pdf_path:
                files_saved["pdf"] = os.path.basename(pdf_path)
            else:
                print("V5 WARNING: PDF conversion failed (LibreOffice not available?)")

        # DOC output option
        if export_format == "doc":
            if update_fn:
                update_fn(task_id, 97, "V5: Convirtiendo a DOC via LibreOffice...")
            doc_path = convert_docx_to_format(docx_output, 'doc', output_folder)
            if doc_path:
                files_saved["doc"] = os.path.basename(doc_path)

        if csv_filename:
            files_saved["csv"] = csv_filename
        if links_docx_filename:
            files_saved["links_docx"] = links_docx_filename

        if update_fn:
            update_fn(task_id, 100, "¡Proceso V5 Linux Completado!")

        return files_saved

    except Exception as e:
        print(f"V5 Fatal compile error (Linux): {e}")
        traceback.print_exc()
        return None


# ============================================
# AI DOCX BUILDER (shared utility)
# ============================================

def build_ai_docx(ai_text, title, output_docx_path):
    """Parse AI text and create a formatted DOCX."""
    ai_doc = Document()
    ai_doc.add_heading(title, level=1)
    for line in ai_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('### '):
            ai_doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('## '):
            ai_doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('# '):
            ai_doc.add_heading(line[2:].strip(), level=1)
        else:
            p = ai_doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
    ai_doc.save(output_docx_path)


def read_document_text(file_path, max_chars=25000):
    """Read text content from a DOCX file for AI analysis."""
    try:
        doc_obj = Document(file_path)
        text_blocks = []
        for el in doc_obj.element.body:
            if el.tag.endswith('p'):
                p_text = Paragraph(el, doc_obj).text.strip()
                if p_text:
                    text_blocks.append(p_text)
            elif el.tag.endswith('tbl'):
                t = Table(el, doc_obj)
                for row in t.rows:
                    row_data = [c.text.replace('\n', ' ').strip() for c in row.cells if c.text.strip()]
                    if row_data:
                        text_blocks.append(" | ".join(row_data))
        return '\n'.join(text_blocks)[:max_chars]
    except Exception:
        return ""
