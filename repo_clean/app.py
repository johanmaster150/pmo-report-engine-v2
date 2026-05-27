import os
import win32com.client
from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
import traceback
import requests
import docx
import re
import zipfile
import shutil

import sys

app = Flask(__name__)
# Configurations
if getattr(sys, 'frozen', False):
    BASE_DIR = os.path.dirname(sys.executable)
    template_folder = os.path.join(sys._MEIPASS, 'templates')
    static_folder = os.path.join(sys._MEIPASS, 'static')
    app.template_folder = template_folder
    app.static_folder = static_folder
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "output")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024 # 500MB max

import pythoncom
import uuid
import threading

global_tasks = {}

def update_progress(task_id, progress, status):
    if task_id and task_id in global_tasks:
        global_tasks[task_id]['progress'] = progress
        global_tasks[task_id]['status'] = status

def generate_executive_summary(project_files, task_id=None):
    """
    Lee una porción del texto de los proyectos, envía a OpenRouter y crea un documento de resumen IA.
    """
    update_progress(task_id, 2, "IA: Iniciando análisis inteligente de proyectos...")
    api_key = os.environ.get("OPENROUTER_API_KEY", "")
    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    combined_text = ""
    
    total = len(project_files)
    for i, file_path in enumerate(project_files):
        update_progress(task_id, 2 + int(10 * (i/total)), f"IA: Leyendo contexto de {os.path.basename(file_path)}...")
        try:
            doc = docx.Document(file_path)
            text_blocks = []
            for el in doc.element.body:
                if el.tag.endswith('p'):
                    p_text = Paragraph(el, doc).text.strip()
                    if p_text: text_blocks.append(p_text)
                elif el.tag.endswith('tbl'):
                    t = Table(el, doc)
                    for row in t.rows:
                        row_data = [c.text.replace('\n', ' ').strip() for c in row.cells if c.text.strip()]
                        if row_data: text_blocks.append(" | ".join(row_data))
            text = "\n".join(text_blocks)
            combined_text += f"\n--- Proyecto: {os.path.basename(file_path)} ---\n"
            combined_text += text[:25000] 
        except:
             pass

    prompt = (
        "Actúa como el Gerente de una PMO (Oficina de Proyectos) para la Alcaldía de Medellín. "
        "A continuación, encontrarás extractos desordenados de reportes de avances de múltiples equipos técnicos este mes. "
        "Tu tarea tiene dos partes:\n"
        "1. Escribir un 'Resumen Ejecutivo Mensual' global, formal y directo resaltando los logros principales consolidados.\n"
        "2. Realizar un resumen específico POR CADA SUBTÍTULO NUMERADO QUE ENCUENTRES (identificados claramente desde 1.1 hasta 1.7, 2.1 hasta 2.3 y 3.1, que suman 11 productos en total) realizando en cada apartado una 'Auditoría Crítica' donde "
        "detectes e indiques explícitamente posibles errores, inconsistencias, datos faltantes, incongruencias, o cosas por mejorar "
        "según el texto provisto para ese numeral específico. Si todo está perfecto en ese numeral, indícalo también, pero sé incisivo.\n\n"
        "Escribe en español de Colombia corporativo. Estructura bien el documento (utilizando los mismos numerales 1.1, 1.2... 3.1 como títulos claros). "
        "NO uses formato Markdown complejo (como asteriscos dobles ** ), usa texto plano bien espaciado.\n\n"
        f"EXTRACTOS:\n{combined_text[:60000]}" # Límite max de contexto final string
    )

    update_progress(task_id, 15, "IA: Enviando datos a OpenRouter (esto toma aprox 15-30s)...")
    try:
        response = requests.post(
            url="https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            },
            json={
                "model": "arcee-ai/trinity-large-preview:free",
                "messages": [{"role": "user", "content": prompt}]
            }
        )
        if response.status_code == 200:
            update_progress(task_id, 35, "IA: Estructurando y Estilizando reporte en Memoria...")
            ai_text = response.json()['choices'][0]['message']['content'].strip()
            
            # Crear un doc temporal con la respuesta
            output_ai_filename_docx = "00_Resumen_y_Auditoria_IA.docx"
            summary_path_docx = os.path.join(OUTPUT_FOLDER, output_ai_filename_docx)
            output_ai_filename_pdf = "00_Resumen_y_Auditoria_IA.pdf"
            summary_path_pdf = os.path.join(OUTPUT_FOLDER, output_ai_filename_pdf)
            
            ai_doc = docx.Document()
            ai_doc.add_heading('Resumen de Gestión y Auditoría (IA)', level=1)
            
            # Parse Markdown Styles
            for line in ai_text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                if line.startswith('### '):
                    ai_doc.add_heading(line[4:].strip(), level=3)
                elif line.startswith('## '):
                    ai_doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith('# '):
                    ai_doc.add_heading(line[2:].strip(), level=1)
                else:
                    p = ai_doc.add_paragraph()
                    # Splitting by bold tags **
                    parts = re.split(r'(\*\*.*?\*\*)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            p.add_run(part[2:-2]).bold = True
                        else:
                            p.add_run(part)

            ai_doc.save(summary_path_docx)
            update_progress(task_id, 38, "IA: Transformando reporte a PDF Premium...")
            
            # Export to PDF via Win32 so it can be viewed on Web
            try:
                pythoncom.CoInitialize()
                word_ai = win32com.client.DispatchEx("Word.Application")
                word_ai.Visible = False
                word_ai.ScreenUpdating = False
                word_ai.DisplayAlerts = False
                w_doc = word_ai.Documents.Open(summary_path_docx)
                w_doc.SaveAs(summary_path_pdf, FileFormat=17) # 17 is wdFormatPDF
                w_doc.Close()
            except Exception as e_pdf:
                print(f"Error convirtiendo IA docx a PDF: {e_pdf}")
            finally:
                try: 
                    word_ai.Quit()
                    pythoncom.CoUninitialize()
                except: pass

            print("¡Auditoría IA completada con éxito!")
            update_progress(task_id, 40, "IA: ¡Auditoría construida! Pasando al ensamblador Win32...")
            
            # Retornamos dict con ambos path
            return {"docx": output_ai_filename_docx, "pdf": output_ai_filename_pdf}
        else:
            err = f"API Request Falló: {response.text}"
            if task_id: global_tasks[task_id]['error_ai'] = err
            print(err)
            return None
    except Exception as e:
        err = f"Error IA Local: {e}"
        if task_id: global_tasks[task_id]['error_ai'] = err
        print(err)
        return None

def compile_unified_word(master_template_path, project_files, output_path, font_family="default", export_format="word", task_id=None):
    """
    Background job that runs Win32 COM locally to compile the word documents.
    """
    update_progress(task_id, 45, "Win32: Cargando Instancia Backend Automática...")
    print("Iniciando MS Word (Fast Mode) en servidor...")
    pythoncom.CoInitialize()
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.ScreenUpdating = False
    word.DisplayAlerts = False
    
    try:
        doc = word.Documents.Open(master_template_path)
        selection = word.Selection
        selection.EndKey(Unit=6) # wdStory = 6
        
        total_p = len(project_files)
        for i, doc_path in enumerate(project_files):
            # Advance progress from 45% -> 70%
            update_progress(task_id, 45 + int(25 * (i/total_p)), f"Win32: Limpiando formato origen de {os.path.basename(doc_path)}...")
            print(f"Insertando (Copy/Paste Inteligente): {os.path.basename(doc_path)}")
            try:
                # 1. Abrimos el documento origen
                src_doc = word.Documents.Open(doc_path, ReadOnly=True, Visible=False)
                # 2. Copiamos su contenido al portapapeles
                src_doc.Content.Copy()
                # 3. Pegamos usando wdUseDestinationStylesRecovery (19). 
                # Esto es el equivalente a "Pegar y hacer coincidir con formato de destino"
                # Limpiando sucios del doc original (HTML backgrounds) y reteniendo imágenes/tablas.
                selection.PasteAndFormat(19)
                src_doc.Close(False)
                
                selection.InsertParagraphAfter()
                selection.EndKey(Unit=6)
            except Exception as e:
                print(f"Error Copy/Paste en {doc_path}: {e}. Intentando fallback...")
                try:
                    if 'src_doc' in locals():
                        src_doc.Close(False)
                except:
                    pass
                # Fallback al viejo InsertFile por si el portapapeles de Windows colapsa
                try:
                    selection.InsertFile(FileName=doc_path)
                    selection.InsertParagraphAfter()
                    selection.EndKey(Unit=6)
                except Exception as e2:
                    print(f"Fallback falló: {e2}")

        update_progress(task_id, 75, "Win32: Aplicando AutoFit masivo a tablas...")
        print("Aplicando AutoFit masivo a tablas...")
        for table in doc.Tables:
            try:
                table.AutoFitBehavior(2) # wdAutoFitWindow
                table.Rows.Alignment = 1 # Center
                table.Borders.Enable = True # Convierte datos sin trazos al formato tabla visible unificado
            except:
                pass

        update_progress(task_id, 85, "Win32: Destruyendo Sombras Grises y formateando textos...")
        print("Aplicando estandarizaciones de estilo (Limpieza de color y fuente)...")
        try:
            # 1. Elimina masivamente todos los comentarios del documento
            doc.DeleteAllComments()
            
            # Selecciona todo el rango y usa Buscar y Reemplazar nativo para borrar Resaltados
            word.Selection.WholeStory()
            find = word.Selection.Find
            find.ClearFormatting()
            find.Highlight = True # Buscar texto que tenga cualquier highlight
            find.Replacement.ClearFormatting()
            find.Replacement.Highlight = False # Sustituir quitando el highlight
            find.Execute(FindText="", ReplaceWith="", Format=True, Replace=2) # 2 = wdReplaceAll
            
            # Borrar Background Shading (sombreado) asegurándose de no romper los colores de las Tablas
            # wdWithInTable = 12
            for p in doc.Paragraphs:
                if not p.Range.Information(12): 
                    # Limpieza agresiva del Highlighting (resaltado amarillo tradicional)
                    p.Range.HighlightColorIndex = 0
                    
                    # NUCLEAR: Pintar el fondo de blanco absoluto (wdColorWhite = 16777215)
                    # Esto sepulta finalmente el problema de "Fondo Gris / Amarillo" heredado de formato web HTML
                    # pero respeta las imagenes, tablas y el color de texto.
                    p.Range.Font.Shading.BackgroundPatternColor = 16777215
                    p.Range.ParagraphFormat.Shading.BackgroundPatternColor = 16777215
                    
                    # Remover texturas adicionales
                    p.Range.Font.Shading.Texture = 0
                    p.Range.ParagraphFormat.Shading.Texture = 0
            
            # 4. Unificar la fuente si el usuario la seleccionó
            if font_family != "default":
                print(f"Cambio global de tipografía a: {font_family}")
                doc.Content.Font.Name = font_family
                
        except Exception as e:
            print(f"Advertencia aplicando estandarizaciones: {e}")

        update_progress(task_id, 98, "Win32: Guardando Ensamblaje Pesado en disco...")
        
        # Decide export formats
        files_saved = {}
        if export_format in ["word", "both"]:
            doc.SaveAs(output_path)
            files_saved["docx"] = os.path.basename(output_path)
        
        if export_format in ["pdf", "both"]:
            update_progress(task_id, 99, "Win32: Imprimiendo Reporte Maestro PDF (Nativo)...")
            pdf_path = output_path.replace(".docx", ".pdf")
            doc.SaveAs(pdf_path, FileFormat=17) # wdFormatPDF
            files_saved["pdf"] = os.path.basename(pdf_path)

        doc.Close()
        update_progress(task_id, 100, "¡Proceso Completado!")
        print("¡Proceso de generación web completado!")
        return files_saved
        
    except Exception as e:
        print(f"Fatal compile error: {e}")
        traceback.print_exc()
        return None
    finally:
        word.ScreenUpdating = True
        word.DisplayAlerts = True
        word.Quit()
        pythoncom.CoUninitialize()

@app.route("/", methods=["GET"])
def index():
    return render_template("index.html")

@app.route("/api/progress/<task_id>")
def get_progress(task_id):
    if task_id in global_tasks:
        return jsonify(global_tasks[task_id])
    return jsonify({"error": "Task not found"}), 404

def worker_compilation(task_id, master_path, project_files, output_path, font_family, generate_ai, mode, export_format):
    try:
        ai_filenames = None
        if generate_ai:
            ai_filenames = generate_executive_summary(project_files, task_id)

        saved_files = compile_unified_word(master_path, project_files, output_path, font_family, export_format, task_id)
        
        if saved_files:
            # We return dicts of downloadUrl paths for word and pdf
            global_tasks[task_id]['downloadUrls'] = {fmt: f"/download/{fname}" for fmt, fname in saved_files.items()}
            
            if ai_filenames:
                global_tasks[task_id]['aiDownloadUrls'] = {fmt: f"/download/{fname}" for fmt, fname in ai_filenames.items()}
                
            global_tasks[task_id]['success'] = True
        else:
            global_tasks[task_id]['error'] = "Falló la compilación del documento con MS Word localmente."
    except Exception as e:
        global_tasks[task_id]['error'] = f"Excepción crítica durante worker: {e}"
        traceback.print_exc()
    finally:
        global_tasks[task_id]['completed'] = True

@app.route("/api/generate", methods=["POST"])
def generate_report():
    print("Received async generation request")
    
    # 0. Generate Task ID first so we can use it for unique filenames
    task_id = str(uuid.uuid4())
    
    global_tasks[task_id] = {
        "progress": 0, "status": "Iniciando colas en servidor...", "completed": False, 
        "success": False, "error": None, "error_ai": None, "downloadUrl": None, "aiDownloadUrl": None
    }

    try:
        # 1. Save Master Template
        if 'masterTemplate' not in request.files:
            return jsonify({"error": "No has subido la plantilla maestra"}), 400
        
        master_file = request.files['masterTemplate']
        master_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{task_id}_master_template.docx')
        master_file.save(master_path)
        
        # 2. Save Projects in Requested Order
        project_files = []
        ordered_keys = request.form.getlist('projectIds[]')
        
        for pk in ordered_keys:
            proj_file = request.files.get(f'projectFile_{pk}')
            if proj_file:
                safe_name = secure_filename(proj_file.filename)
                save_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{task_id}_{safe_name}')
                proj_file.save(save_path)
                project_files.append(save_path)

        if not project_files:
            return jsonify({"error": "No se subieron archivos de proyecto válidos"}), 400

        # 3. Extraer la tipografia, modo y formato de exportacion
        font_family = request.form.get("fontFamily", "default")
        mode = request.form.get("mode", "unified")
        generate_ai = request.form.get("generateAI") == "true"
        export_format = request.form.get("exportFormat", "word")

        if mode == "split":
            template2_file = request.files.get("template2")
            if template2_file:
                t2_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{task_id}_master_template_pago2.docx')
                template2_file.save(t2_path)
                print("INFO: Modo Pagos Divididos seleccionado. (En desarrollo)")

        # 4. Iniciar hilo en background para compilación y dar taskId a frontend
        output_filename = f"Reporte_Mensual_Generado_{task_id[:8]}.docx"
        output_path = os.path.join(OUTPUT_FOLDER, output_filename)
        
        thread = threading.Thread(target=worker_compilation, args=(task_id, master_path, project_files, output_path, font_family, generate_ai, mode, export_format))
        thread.start()
        
        return jsonify({"taskId": task_id, "status": "Task Started"})

    except Exception as e:
        print(f"Error procesando peticion POST V1: {e}")
        return jsonify({"error": f"Error al procesar archivos: {str(e)}"}), 500

@app.route("/download/<filename>")
def download_file(filename):
    path = os.path.join(OUTPUT_FOLDER, filename)
    return send_file(path, as_attachment=True)

# ----------------- V2 PARALLEL PROCESS -----------------

import csv

def check_cti_path(url):
    """Check if a URL belongs to the CTI strategic projects path."""
    from urllib.parse import unquote
    decoded_url = unquote(url).lower()
    cti_indicators = [
        'proyectos estratégicos cti', 'proyectos estrategicos cti',
        'proyectos%20estrat', '03productos', '03 productos',
        'documentos'
    ]
    # Must match at least CTI + one more indicator for confidence
    matches = sum(1 for ind in cti_indicators if ind in decoded_url)
    if matches >= 2:
        return "✅ Sí - En ruta CTI"
    elif matches == 1:
        return "⚠️ Posible - Ruta parcial CTI"
    return "❌ No"

def extract_and_download_links(word, doc, output_dir, task_id=None, product_sections=None):
    update_progress(task_id, 90, "Win32 V2: Extrayendo enlaces y evidencias...")
    evidencias_dir = os.path.join(output_dir, "Evidencias_Descargadas")
    os.makedirs(evidencias_dir, exist_ok=True)
    
    links_data = []  # Each row: [numeracion, url_pura, texto_documento, en_ruta_cti, estado]
    
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    links_doc = docx.Document()
    links_doc.add_heading('Listado de Enlaces y Evidencias', level=1)
    
    if product_sections is None:
        product_sections = []
        
    try:
        hyperlinks = doc.Hyperlinks
        total_links = hyperlinks.Count
        link_counters = {}
        
        for i in range(1, total_links + 1):
            hl = hyperlinks.Item(i)
            url = getattr(hl, "Address", None)
            text_to_display = getattr(hl, "TextToDisplay", str(url))
            if not url or not url.startswith("http"):
                continue
            
            hl_start = hl.Range.Start
            current_prod = None
            for sec in reversed(product_sections):
                if hl_start >= (sec["start"] - 150):
                    current_prod = sec["prod_id"]
                    break
                    
            if current_prod:
                link_counters[current_prod] = link_counters.get(current_prod, 0) + 1
                seq_id = f"{current_prod}.{link_counters[current_prod]}"
            else:
                link_counters["general"] = link_counters.get("general", 0) + 1
                seq_id = f"G.{link_counters['general']}"
            
            # Check if URL is in CTI path
            cti_status = check_cti_path(url)
            
            status = "Fallido (Requiere Autenticación)"
            try:
                res = requests.get(url, timeout=5, stream=True)
                if res.status_code == 200 and 'text/html' not in res.headers.get('Content-Type', ''):
                    filename = secure_filename(url.split('/')[-1])
                    if not filename:
                        filename = f"evidencia_{i}.file"
                    filepath = os.path.join(evidencias_dir, filename)
                    with open(filepath, 'wb') as f:
                        for chunk in res.iter_content(chunk_size=8192):
                            if chunk:
                                f.write(chunk)
                    status = "Descargado"
            except Exception as e:
                pass
            
            # 5-column data: [#, URL pura, Texto en doc, ¿En ruta CTI?, Estado]
            links_data.append([seq_id, url, text_to_display, cti_status, status])

        csv_filename = None
        docx_filename = None
        if links_data:
            csv_filename = "00_Reporte_Enlaces_Evidencias.csv"
            docx_filename = "02_Listado_Enlaces_y_Evidencias.docx"
            csv_path = os.path.join(output_dir, csv_filename)
            docx_path = os.path.join(output_dir, docx_filename)
            
            # --- CSV with 5 columns ---
            with open(csv_path, 'w', newline='', encoding='utf-8-sig') as f:
                writer = csv.writer(f)
                writer.writerow(["#", "Enlace (URL)", "Texto en Documento", "¿En ruta CTI?", "Estado de Descarga"])
                writer.writerows(links_data)
            
            # --- DOCX with formatted table ---
            table = links_doc.add_table(rows=1, cols=5)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'
            
            # Header row
            headers = ["#", "Enlace (URL)", "Texto en Documento", "¿En ruta CTI?", "Estado"]
            for col_idx, header_text in enumerate(headers):
                cell = table.rows[0].cells[col_idx]
                cell.text = header_text
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(255, 255, 255)
                from docx.oxml.ns import qn
                shading_elm = cell._element.get_or_add_tcPr()
                shading = shading_elm.makeelement(qn('w:shd'), {
                    qn('w:fill'): '2F5496',
                    qn('w:val'): 'clear'
                })
                shading_elm.append(shading)
            
            # Data rows
            for row_data in links_data:
                row_cells = table.add_row().cells
                for col_idx, value in enumerate(row_data):
                    row_cells[col_idx].text = str(value)
                    for paragraph in row_cells[col_idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
            
            # Set column widths
            try:
                for row in table.rows:
                    row.cells[0].width = Inches(0.6)   # #
                    row.cells[1].width = Inches(3.0)    # URL
                    row.cells[2].width = Inches(2.0)    # Texto
                    row.cells[3].width = Inches(1.2)    # CTI
                    row.cells[4].width = Inches(1.0)    # Estado
            except:
                pass
            
            links_doc.save(docx_path)
            return csv_filename, docx_filename
    except Exception as e:
        print("Error extracting links in V2:", e)
    return None, None

def compile_unified_word_v2(master_template_path, project_files, output_path, font_family="default", export_format="word", task_id=None):
    update_progress(task_id, 45, "Win32 V2: Iniciando motor de ensamble nativo (InsertFile)...")
    pythoncom.CoInitialize()
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.ScreenUpdating = False
    word.DisplayAlerts = False
    
    # -- Master list of TOC titles to search for, with their indentation level --
    # Level 1 = bold main headings, Level 2 = sub-sections (Products are extracted dynamically)
    TOC_TITLES = [
        (1, "Informe Mensual"),
        (2, "Programa: Innovación como motor del desarrollo económico del Distrito"),
        (1, "Generalidades"),
        (2, "Objetivo General"),
        (2, "Proyectos por desarrollar"),
        (2, "Descripción informe"),
        (1, "Avance técnico Proyecto: Fortalecimiento del ecosistema CTi"),
        (2, "Avance Financiero y presupuestal"),
        (1, "Avance técnico Desarrollo de negocios de alto impacto"),
        (2, "Avance Financiero y presupuestal"),
        (1, "Avance técnico Proyecto: Centro para la Cuarta Revolución Industrial (C4RI)"),
        (1, "Avance técnico Proyecto: Centro para la Cuarta Revolución Industrial"),
        (2, "Avance Financiero y presupuestal"),
    ]
    
    try:
        doc = word.Documents.Open(master_template_path)
        
        # --- STEP 0: Clear master template BODY content (keep format, styles, headers, footers) ---
        update_progress(task_id, 47, "Win32 V2: Extrayendo formato del maestro (sin contenido)...")
        try:
            # Delete ALL existing TOCs from master template
            while doc.TablesOfContents.Count > 0:
                doc.TablesOfContents(1).Delete()
            # Clear ALL body text content, preserving section formatting
            doc.Content.Delete()
            print("Master template content cleared. Format/styles/headers/footers preserved.")
        except Exception as e:
            print(f"Warning clearing master content: {e}")
        
        selection = word.Selection
        selection.EndKey(Unit=6)  # wdStory = 6
        
        # --- STEP 1: Insert project files in order ---
        total_p = len(project_files)
        for i, doc_path in enumerate(project_files):
            update_progress(task_id, 48 + int(25 * (i/total_p)), f"Win32 V2: Insertando estructuralmente {os.path.basename(doc_path)}...")
            print(f"V2 Insertando nativo: {os.path.basename(doc_path)}")
            try:
                if i > 0:
                    # Section break between files, but not before the first one
                    selection.InsertBreak(Type=2)  # wdSectionBreakNextPage
                selection.InsertFile(FileName=doc_path)
                selection.EndKey(Unit=6)
            except Exception as e:
                print(f"Error V2 InsertFile en {doc_path}: {e}")

        # --- STEP 2: Apply unified font ---
        if font_family != "default":
            update_progress(task_id, 74, "Win32 V2: Aplicando tipografía unificada...")
            try:
                doc.Content.Font.Name = font_family
            except: pass

        # --- STEP 3 & 3.5: AutoFit, Bordes Visibles, y Renumeración de Productos ---
        update_progress(task_id, 75, "Win32 V2: Aplicando bordes, extrayendo productos y renumerando...")
        try:
            dynamic_products_toc = []
            product_sections = []  # list of {start, prod_id}

            prefix_pattern = re.compile(r'^\s*(\d+\.\d+)\b')
            
            # ===== KNOWN SECTION IDS =====
            KNOWN_PROD_IDS = [
                '1.1', '1.2', '1.3', '1.4', '1.5', '1.6', '1.7',
                '2.1', '2.2', '2.3', '2.4',
                '3.1'
            ]

            # ===== Apply AutoFit and Borders to ALL tables first =====
            all_tables_info = []  # list of (table_obj, table_range_start, table_range_end)
            for table_idx in range(1, doc.Tables.Count + 1):
                table = doc.Tables(table_idx)
                try: table.AutoFitBehavior(2)
                except: pass
                try: table.Rows.Alignment = 1
                except: pass
                try: table.Borders.Enable = True
                except: pass
                try: table.Borders.InsideLineStyle = 1
                except: pass
                try: table.Borders.OutsideLineStyle = 1
                except: pass
                try:
                    all_tables_info.append((table, table.Range.Start, table.Range.End))
                except:
                    pass

            # ===== STRATEGY A: Header-keyword detection (EXPANDED) =====
            for table, t_start, t_end in all_tables_info:
                try:
                    num_rows = table.Rows.Count
                    num_cols = table.Columns.Count

                    # Expanded header check: much broader keywords
                    header_looks_like_product = False
                    if num_cols >= 2:
                        h1 = table.Cell(1,1).Range.Text.replace('\r\a', '').strip().lower()
                        h2 = table.Cell(1,2).Range.Text.replace('\r\a', '').strip().lower()
                        # Check additional columns if available
                        h3 = ''
                        if num_cols >= 3:
                            try: h3 = table.Cell(1,3).Range.Text.replace('\r\a', '').strip().lower()
                            except: pass
                        
                        all_headers = h1 + ' ' + h2 + ' ' + h3
                        product_keywords = [
                            '#', 'n°', 'no.', 'num', 'producto', 'nombre', 'descripcion',
                            'indicador', 'meta', 'actividad', 'avance', 'resultado',
                            'ponderación', 'ponderacion', 'componente', 'hito',
                            'entregable', 'evidencia', 'porcentaje', 'cumplimiento',
                            'línea base', 'linea base'
                        ]
                        if any(kw in all_headers for kw in product_keywords):
                            header_looks_like_product = True

                    if not header_looks_like_product:
                        continue

                    # Scan rows (up to 10 now) to find a prod_id
                    prod_id = None
                    prod_name = None
                    for row_idx in range(2, min(num_rows + 1, 11)):  # Up to 10 data rows
                        try:
                            val1 = table.Cell(row_idx, 1).Range.Text.replace('\r\a', '').strip()
                            val2 = table.Cell(row_idx, 2).Range.Text.replace('\r\a', '').strip() if num_cols >= 2 else ''
                        except:
                            continue

                        m1 = prefix_pattern.match(val1)
                        m2 = prefix_pattern.match(val2)

                        if m1 and val2 and not prefix_pattern.match(val2):
                            prod_id = m1.group(1)
                            prod_name = val2.strip()
                            break
                        elif m2 and val1 and not prefix_pattern.match(val1):
                            prod_id = m2.group(1)
                            prod_name = val1.strip()
                            break
                        elif m1:
                            prod_id = m1.group(1)
                            prod_name = val2 if val2 else val1
                            break

                    if prod_id and prod_name and len(prod_name) > 2:
                        existing_ids = [s['prod_id'] for s in product_sections]
                        if prod_id not in existing_ids:
                            table.Range.InsertParagraphBefore()
                            p_rng = doc.Range(table.Range.Start - 1, table.Range.Start - 1)
                            p_rng.Text = f"Tabla de Producto: {prod_name}"

                            dynamic_products_toc.append({
                                "prod_id": prod_id,
                                "prod_name": prod_name,
                                "table": table
                            })
                            product_sections.append({
                                "start": table.Range.Start,
                                "prod_id": prod_id
                            })

                except Exception as e_tbl:
                    print(f"  Strategy A table detection error: {e_tbl}")
                    pass

            print(f"Strategy A found {len(product_sections)} product sections via header keywords.")

            # ===== STRATEGY B: Search for KNOWN section IDs in paragraphs, then find nearest table =====
            # Rebuild all_tables_info since Strategy A may have shifted positions
            all_tables_info = []
            for table_idx in range(1, doc.Tables.Count + 1):
                try:
                    table = doc.Tables(table_idx)
                    all_tables_info.append((table, table.Range.Start, table.Range.End))
                except:
                    pass

            existing_ids_set = set(s['prod_id'] for s in product_sections)

            for known_id in KNOWN_PROD_IDS:
                if known_id in existing_ids_set:
                    continue  # Already found by Strategy A

                # Search the entire document for a paragraph that starts with or contains this ID
                search_range = doc.Content.Duplicate
                search_range.Find.ClearFormatting()
                
                # Search for the known section ID pattern (e.g. "1.1" or "2.1" in text)
                found = search_range.Find.Execute(
                    FindText=known_id,
                    MatchCase=False,
                    MatchWholeWord=False,
                    Forward=True,
                    Wrap=0
                )

                if found:
                    found_pos = search_range.Start
                    # Verify it's actually a section heading: the paragraph text should start with or prominently contain the ID
                    try:
                        found_para_text = search_range.Paragraphs(1).Range.Text.strip()
                    except:
                        found_para_text = ""
                    
                    # Validate: the found text should be in a context that looks like a section header
                    # (not just any random occurrence of "1.1" in a table cell)
                    is_section_heading = False
                    check_text = found_para_text.replace('\r\a', '').strip()
                    if check_text:
                        # Must start with the ID or be the ID followed by descriptive text
                        id_pattern = re.compile(r'^\s*' + re.escape(known_id) + r'\s*\.?\s*\w', re.IGNORECASE)
                        id_pattern_short = re.compile(r'^\s*' + re.escape(known_id) + r'\s*$')
                        prod_keyword_pattern = re.compile(r'(?:producto|product|indicador|meta|actividad|componente)', re.IGNORECASE)
                        if id_pattern.match(check_text) or id_pattern_short.match(check_text):
                            is_section_heading = True
                        elif known_id in check_text and prod_keyword_pattern.search(check_text):
                            is_section_heading = True

                    if not is_section_heading:
                        # Try searching inside tables directly for this ID
                        pass

                    # Find the nearest table AFTER this position
                    nearest_table = None
                    nearest_dist = float('inf')
                    for tbl, t_start, t_end in all_tables_info:
                        if t_start >= found_pos and (t_start - found_pos) < nearest_dist:
                            # Make sure this table isn't already claimed by another product
                            already_claimed = False
                            for sec in product_sections:
                                if abs(sec['start'] - t_start) < 50:
                                    already_claimed = True
                                    break
                            if not already_claimed:
                                nearest_dist = t_start - found_pos
                                nearest_table = tbl

                    if nearest_table and nearest_dist < 5000:  # Must be within ~5000 chars
                        # Try to extract a product name from the table or nearby text
                        prod_name_b = None
                        try:
                            nr = nearest_table.Rows.Count
                            nc = nearest_table.Columns.Count
                            for row_idx in range(1, min(nr + 1, 11)):
                                try:
                                    for col_idx in range(1, min(nc + 1, 5)):
                                        cell_text = nearest_table.Cell(row_idx, col_idx).Range.Text.replace('\r\a', '').strip()
                                        if cell_text and len(cell_text) > 5 and not prefix_pattern.match(cell_text):
                                            if not any(kw in cell_text.lower() for kw in ['#', 'n°', 'indicador', 'meta', 'actividad']):
                                                prod_name_b = cell_text[:100]
                                                break
                                    if prod_name_b:
                                        break
                                except:
                                    continue
                        except:
                            pass

                        if not prod_name_b:
                            prod_name_b = f"Producto {known_id}"

                        # Insert label paragraph before table
                        try:
                            nearest_table.Range.InsertParagraphBefore()
                            p_rng = doc.Range(nearest_table.Range.Start - 1, nearest_table.Range.Start - 1)
                            p_rng.Text = f"Tabla de Producto: {prod_name_b}"
                        except:
                            pass

                        dynamic_products_toc.append({
                            "prod_id": known_id,
                            "prod_name": prod_name_b,
                            "table": nearest_table
                        })
                        product_sections.append({
                            "start": nearest_table.Range.Start,
                            "prod_id": known_id
                        })
                        existing_ids_set.add(known_id)
                        print(f"  Strategy B: Found section '{known_id}' -> table at pos {nearest_table.Range.Start}")

            print(f"Strategy B total: now {len(product_sections)} product sections (including both strategies).")

            # Sort product_sections by document position
            product_sections.sort(key=lambda x: x['start'])
            doc_end = doc.Content.End

            # Build list of (prod_id, section_start, section_end) for position-based lookup
            section_ranges = []
            for i, sec in enumerate(product_sections):
                s_start = sec['start']
                # section ends where the next one begins (or at doc end)
                s_end = product_sections[i + 1]['start'] if i + 1 < len(product_sections) else doc_end
                section_ranges.append((sec['prod_id'], s_start, s_end))

            def get_prod_id_for_pos(pos):
                """Return prod_id for a given character position, or None."""
                for pid, s, e in section_ranges:
                    if s <= pos < e:
                        return pid
                return None

            # --- Renumber 'Tabla ...' captions using position ranges ---
            table_seq_by_prefix = {}  # {prod_id: seq_counter}
            total_paras = doc.Paragraphs.Count

            for p_idx in range(1, total_paras + 1):
                try:
                    para = doc.Paragraphs(p_idx)
                    text = para.Range.Text.strip()
                    if not text:
                        continue

                    if re.match(r'^\s*(?:[\d\.]+\s*)?tabla\b', text, re.IGNORECASE):
                        para_pos = para.Range.Start
                        current_prefix = get_prod_id_for_pos(para_pos)

                        if current_prefix:
                            seq = table_seq_by_prefix.get(current_prefix, 0) + 1
                            table_seq_by_prefix[current_prefix] = seq
                            new_text = re.sub(
                                r'^\s*(?:[\d\.]+\s*)?tabla\b',
                                f'{current_prefix}.{seq} Tabla',
                                text, flags=re.IGNORECASE
                            )
                        else:
                            new_text = re.sub(
                                r'^\s*(?:[\d\.]+\s*)?tabla\b',
                                'Tabla',
                                text, flags=re.IGNORECASE
                            )

                        rng = para.Range
                        rng.Find.ClearFormatting()
                        rng.Find.Replacement.ClearFormatting()
                        rng.Find.Execute(
                            FindText=text, ReplaceWith=new_text, Replace=1,
                            MatchCase=False, MatchWholeWord=False
                        )
                except:
                    pass

            print(f"Table detection found {len(product_sections)} product sections. Renumbering complete.")
        except Exception as e:
            print(f"Error during table processing: {e}")

        # --- STEP 4: Empty Space Optimization ---
        update_progress(task_id, 80, "Win32 V2: Optimizando espacios vacíos...")
        try:
            find_rng = doc.Content
            for _ in range(5):
                find_rng.Find.ClearFormatting()
                find_rng.Find.Replacement.ClearFormatting()
                replaced = find_rng.Find.Execute(
                    FindText="^p^p^p",
                    ReplaceWith="^p^p",
                    Replace=2,  # wdReplaceAll
                    Forward=True,
                    Wrap=1  # wdFindContinue
                )
                if not replaced:
                    break
            print("Empty space optimization complete.")
        except Exception as e:
            print(f"Error during empty space optimization: {e}")

        # --- STEP 5: Custom TOC Builder ---
        update_progress(task_id, 83, "Win32 V2: Construyendo Índice personalizado con títulos encontrados...")
        try:
            toc_entries = []
            avance_financiero_count = 0  
            
            for level, title_search in TOC_TITLES:
                search_range = doc.Content.Duplicate
                search_range.Find.ClearFormatting()
                found = search_range.Find.Execute(
                    FindText=title_search,
                    MatchCase=False,
                    MatchWholeWord=False,
                    Forward=True,
                    Wrap=0  
                )
                
                if found:
                    page_num = search_range.Information(3)  
                    
                    if "Avance Financiero y presupuestal" in title_search:
                        avance_financiero_count += 1
                        if avance_financiero_count > 1:
                            for _ in range(avance_financiero_count - 1):
                                search_range.Collapse(0)  
                                found_again = search_range.Find.Execute(
                                    FindText=title_search,
                                    MatchCase=False,
                                    MatchWholeWord=False,
                                    Forward=True,
                                    Wrap=0
                                )
                                if found_again:
                                    page_num = search_range.Information(3)
                                else:
                                    found = False
                                    break
                    
                    if found:
                        bookmark_name = f"_TOC_Custom_{len(toc_entries)}"
                        try:
                            doc.Bookmarks.Add(bookmark_name, search_range)
                        except:
                            bookmark_name = None
                        
                        toc_entries.append({
                            "level": level,
                            "title": title_search,
                            "page": page_num,
                            "bookmark": bookmark_name,
                            "start": search_range.Start
                        })
                        print(f"  TOC: Found '{title_search}' on page {page_num}")
                else:
                    print(f"  TOC: NOT found '{title_search}' - skipping")

            for prod in dynamic_products_toc:
                try:
                    page_num = prod["table"].Range.Information(3)
                    bookmark_name = f"_TOC_Custom_{len(toc_entries)}"
                    try:
                        p = prod["table"].Range.Previous(4, 1)
                        if not p: p = prod["table"].Range
                        doc.Bookmarks.Add(bookmark_name, p)
                    except:
                        bookmark_name = None
                    
                    toc_entries.append({
                        "level": 2,
                        "title": f"Producto {prod['prod_id']}: {prod['prod_name']}",
                        "page": page_num,
                        "bookmark": bookmark_name,
                        "start": prod["table"].Range.Start
                    })
                except Exception as e:
                    print("Error dynamic TOC product:", e)

            toc_entries.sort(key=lambda x: x["start"])
            
            # Build and insert TOC - find placeholder or insert at start
            if toc_entries:
                update_progress(task_id, 86, "Win32 V2: Insertando tabla de contenido personalizada...")

                # Build complete TOC lines
                toc_lines = []
                for entry in toc_entries:
                    indent = "    " if entry["level"] == 2 else ""
                    title_display = entry["title"]
                    page_str = str(entry["page"])
                    dots = '.' * max(3, 80 - len(indent) - len(title_display) - len(page_str))
                    toc_lines.append((entry["level"], f"{indent}{title_display} {dots} {page_str}"))

                # Search for placeholder paragraph that contains 'Tabla de Contenido' (the heading)
                toc_insert_para_idx = None
                toc_placeholder_range = None
                total_paras_now = doc.Paragraphs.Count
                for p_idx in range(1, total_paras_now + 1):
                    try:
                        p_text = doc.Paragraphs(p_idx).Range.Text.strip()
                        if re.search(r'tabla\s+de\s+contenido', p_text, re.IGNORECASE):
                            toc_insert_para_idx = p_idx
                            toc_placeholder_range = doc.Paragraphs(p_idx).Range
                            break
                    except:
                        pass

                if toc_placeholder_range:
                    # Replace the placeholder heading's paragraph content and insert entries after it
                    # Format existing heading paragraph
                    toc_placeholder_range.Font.Name = "Arial"
                    toc_placeholder_range.Font.Size = 16
                    toc_placeholder_range.Font.Bold = True
                    toc_placeholder_range.ParagraphFormat.Alignment = 1  # Center

                    # Insert TOC lines AFTER the heading paragraph
                    # We build the full block as a single InsertAfter to keep it together
                    toc_block = "\r\n"
                    for _, line in toc_lines:
                        toc_block += line + "\r\n"
                    toc_block += "\r\n"

                    # Collapse to end of heading paragraph and insert
                    insert_pos = toc_placeholder_range.Duplicate
                    insert_pos.Collapse(0)  # 0 = wdCollapseEnd
                    insert_pos.InsertAfter(toc_block)

                    # Now format the newly inserted paragraphs
                    # They follow toc_insert_para_idx + 1
                    for i, (level, _) in enumerate(toc_lines):
                        para_idx = toc_insert_para_idx + 2 + i  # +1 for blank, +1 for 1-based
                        try:
                            para = doc.Paragraphs(para_idx)
                            para.Range.Font.Name = "Arial"
                            para.Range.Font.Bold = (level == 1)
                            para.Range.Font.Size = 11 if level == 1 else 10
                            para.Range.ParagraphFormat.Alignment = 0
                        except:
                            pass

                    print(f"Custom TOC inserted at placeholder (para {toc_insert_para_idx}) with {len(toc_entries)} entries.")

                else:
                    # Fallback: prepend at document start if no placeholder found
                    toc_text = "Tabla de Contenido\r\n\r\n"
                    for _, line in toc_lines:
                        toc_text += line + "\r\n"
                    toc_text += "\r\n"

                    start_range = doc.Range(Start=doc.Content.Start, End=doc.Content.Start)
                    start_range.InsertBefore(toc_text)

                    first_para = doc.Paragraphs(1)
                    first_para.Range.Font.Name = "Arial"
                    first_para.Range.Font.Size = 16
                    first_para.Range.Font.Bold = True
                    first_para.Range.ParagraphFormat.Alignment = 1

                    for i, (level, _) in enumerate(toc_lines):
                        para_idx = i + 3
                        try:
                            para = doc.Paragraphs(para_idx)
                            para.Range.Font.Name = "Arial"
                            para.Range.Font.Bold = (level == 1)
                            para.Range.Font.Size = 11 if level == 1 else 10
                            para.Range.ParagraphFormat.Alignment = 0
                        except:
                            pass

                    print(f"Custom TOC prepended at start with {len(toc_entries)} entries (no placeholder found).")
            else:
                print("No TOC entries found in document.")
                
        except Exception as e:
            print(f"Error building custom TOC: {e}")
            traceback.print_exc()

        # Requirement 4: Links extraction
        csv_filename, links_docx_filename = extract_and_download_links(word, doc, OUTPUT_FOLDER, task_id, product_sections)

        update_progress(task_id, 95, "Win32 V2: Guardando Ensamblaje en disco...")
        
        files_saved = {}
        if export_format in ["word", "both"]:
            doc.SaveAs(output_path)
            files_saved["docx"] = os.path.basename(output_path)
            
        if export_format in ["pdf", "both"]:
            update_progress(task_id, 98, "Win32 V2: Exportando PDF...")
            pdf_path = output_path.replace(".docx", ".pdf")
            doc.SaveAs(pdf_path, FileFormat=17) # wdFormatPDF
            files_saved["pdf"] = os.path.basename(pdf_path)
            
        if csv_filename:
            files_saved["csv"] = csv_filename
        if links_docx_filename:
            files_saved["links_docx"] = links_docx_filename

        doc.Close()
        update_progress(task_id, 100, "¡Proceso V2 Completado!")
        return files_saved
        
    except Exception as e:
        print(f"Fatal compile error V2: {e}")
        traceback.print_exc()
        return None
    finally:
        word.ScreenUpdating = True
        word.DisplayAlerts = True
        word.Quit()
        pythoncom.CoUninitialize()

def generate_predictive_report(project_files, task_id=None):
    """
    Analyzes progress per product and generates predictions/suggestions for next month's report.
    """
    update_progress(task_id, 20, "IA Predictiva: Analizando avances por producto...")
    api_key = os.environ.get("OPENROUTER_API_KEY", "")
    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    combined_text = ""
    
    total = len(project_files)
    for i, file_path in enumerate(project_files):
        update_progress(task_id, 20 + int(5 * (i/total)), f"IA Predictiva: Leyendo {os.path.basename(file_path)}...")
        try:
            doc = docx.Document(file_path)
            text_blocks = []
            for el in doc.element.body:
                if el.tag.endswith('p'):
                    p_text = Paragraph(el, doc).text.strip()
                    if p_text: text_blocks.append(p_text)
                elif el.tag.endswith('tbl'):
                    t = Table(el, doc)
                    for row in t.rows:
                        row_data = [c.text.replace('\n', ' ').strip() for c in row.cells if c.text.strip()]
                        if row_data: text_blocks.append(" | ".join(row_data))
            text = "\n".join(text_blocks)
            combined_text += f"\n--- Proyecto: {os.path.basename(file_path)} ---\n"
            combined_text += text[:25000]
        except:
            pass

    prompt = (
        "Actúa como un consultor senior de PMO (Project Management Office) con experiencia en proyectos "
        "de innovación y tecnología para la Alcaldía de Medellín. "
        "A continuación encontrarás extractos de reportes mensuales de avance de múltiples equipos técnicos.\n\n"
        "Tu tarea es generar un INFORME PREDICTIVO Y DE RECOMENDACIONES con la siguiente estructura:\n\n"
        "1. ANÁLISIS DE AVANCE POR SUBTÍTULO O PRODUCTO NUMERADO: Para CADA sección numerada encontrada (identificados claramente desde 1.1 hasta 1.7, 2.1 hasta 2.3 y 3.1, para un total de 11 productos), analiza separadamente:\n"
        "   - Nivel de avance estimado para esa sección específica (porcentaje y justificación)\n"
        "   - Hitos logrados vs pendientes de esa sección\n"
        "   - Riesgos identificados de esa sección\n\n"
        "2. PREDICCIONES PARA EL PRÓXIMO MES: Para cada producto o subtítulo numerado, sugiere:\n"
        "   - Qué resultados concretos deberían poder reportar el mes siguiente basado en la tendencia actual\n"
        "   - Qué datos o evidencias deberían incluir\n"
        "   - Alertas tempranas si hay riesgo de incumplimiento\n\n"
        "3. RECOMENDACIONES DE REDACCIÓN: Para cada subtítulo numerado o equipo, sugiere:\n"
        "   - Métricas clave que deberían destacar\n"
        "   - Información faltante que deberían completar\n"
        "   - Mejoras en la presentación de sus avances\n\n"
        "Escribe en español de Colombia corporativo. Estructura bien el documento respetando los identificadores numéricos (1.1, 1.2...) como títulos claros. "
        "NO uses formato Markdown complejo (como asteriscos dobles ** ), usa texto plano bien espaciado.\n\n"
        f"EXTRACTOS DE REPORTES:\n{combined_text[:60000]}"
    )

    update_progress(task_id, 27, "IA Predictiva: Enviando análisis a OpenRouter...")
    try:
        response = requests.post(
            url="https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            },
            json={
                "model": "arcee-ai/trinity-large-preview:free",
                "messages": [{"role": "user", "content": prompt}]
            }
        )
        if response.status_code == 200:
            update_progress(task_id, 33, "IA Predictiva: Construyendo informe predictivo...")
            ai_text = response.json()['choices'][0]['message']['content'].strip()
            
            output_filename_docx = "01_Informe_Predictivo_IA.docx"
            pred_path_docx = os.path.join(OUTPUT_FOLDER, output_filename_docx)
            output_filename_pdf = "01_Informe_Predictivo_IA.pdf"
            pred_path_pdf = os.path.join(OUTPUT_FOLDER, output_filename_pdf)
            
            pred_doc = docx.Document()
            pred_doc.add_heading('Informe Predictivo y Recomendaciones (IA)', level=1)
            
            for line in ai_text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                if line.startswith('### '):
                    pred_doc.add_heading(line[4:].strip(), level=3)
                elif line.startswith('## '):
                    pred_doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith('# '):
                    pred_doc.add_heading(line[2:].strip(), level=1)
                else:
                    p = pred_doc.add_paragraph()
                    parts = re.split(r'(\*\*.*?\*\*)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            p.add_run(part[2:-2]).bold = True
                        else:
                            p.add_run(part)

            pred_doc.save(pred_path_docx)
            update_progress(task_id, 36, "IA Predictiva: Convirtiendo a PDF...")
            
            try:
                pythoncom.CoInitialize()
                word_pred = win32com.client.DispatchEx("Word.Application")
                word_pred.Visible = False
                word_pred.ScreenUpdating = False
                word_pred.DisplayAlerts = False
                w_doc = word_pred.Documents.Open(pred_path_docx)
                w_doc.SaveAs(pred_path_pdf, FileFormat=17)
                w_doc.Close()
            except Exception as e_pdf:
                print(f"Error convirtiendo Predictivo docx a PDF: {e_pdf}")
            finally:
                try:
                    word_pred.Quit()
                    pythoncom.CoUninitialize()
                except: pass

            print("¡Informe Predictivo IA completado!")
            update_progress(task_id, 38, "IA Predictiva: ¡Informe construido!")
            return {"docx": output_filename_docx, "pdf": output_filename_pdf}
        else:
            err = f"API Request Predictivo Falló: {response.text}"
            print(err)
            return None
    except Exception as e:
        err = f"Error IA Predictiva: {e}"
        print(err)
        return None

def worker_compilation_v2(task_id, master_path, project_files, output_path, font_family, generate_ai, mode, export_format):
    try:
        ai_filenames = None
        ai_predictive_filenames = None
        if generate_ai:
            ai_filenames = generate_executive_summary(project_files, task_id)
            ai_predictive_filenames = generate_predictive_report(project_files, task_id)

        saved_files = compile_unified_word_v2(master_path, project_files, output_path, font_family, export_format, task_id)
        
        if saved_files:
            download_urls = {}
            if "docx" in saved_files: download_urls["docx"] = f"/download/{saved_files['docx']}"
            if "pdf" in saved_files: download_urls["pdf"] = f"/download/{saved_files['pdf']}"
            if "csv" in saved_files: download_urls["csv"] = f"/download/{saved_files['csv']}"
            if "links_docx" in saved_files: download_urls["links_docx"] = f"/download/{saved_files['links_docx']}"
                
            global_tasks[task_id]['downloadUrls'] = download_urls
            
            if ai_filenames:
                global_tasks[task_id]['aiDownloadUrls'] = {fmt: f"/download/{fname}" for fmt, fname in ai_filenames.items()}
            
            if ai_predictive_filenames:
                global_tasks[task_id]['aiPredictiveUrls'] = {fmt: f"/download/{fname}" for fmt, fname in ai_predictive_filenames.items()}
                
            global_tasks[task_id]['success'] = True
        else:
            global_tasks[task_id]['error'] = "Falló la compilación del documento con MS Word localmente."
    except Exception as e:
        global_tasks[task_id]['error'] = f"Excepción crítica durante worker v2: {e}"
        traceback.print_exc()
    finally:
        global_tasks[task_id]['completed'] = True

@app.route("/v2", methods=["GET"])
def index_v2():
    return render_template("index_v2.html")

@app.route("/api/generate_v2", methods=["POST"])
def generate_report_v2():
    print("Received async generation request V2")
    
    # 0. Generate Task ID first so we can use it for unique filenames
    task_id = str(uuid.uuid4())
    
    global_tasks[task_id] = {
        "progress": 0, "status": "Iniciando colas en servidor V2...", "completed": False, 
        "success": False, "error": None, "error_ai": None, "downloadUrl": None, "aiDownloadUrl": None
    }

    try:
        if 'masterTemplate' not in request.files:
            return jsonify({"error": "No has subido la plantilla maestra"}), 400
        
        master_file = request.files['masterTemplate']
        master_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{task_id}_master_template.docx')
        master_file.save(master_path)
        
        project_files = []
        ordered_keys = request.form.getlist('projectIds[]')
        for pk in ordered_keys:
            proj_file = request.files.get(f'projectFile_{pk}')
            if proj_file:
                safe_name = secure_filename(proj_file.filename)
                save_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{task_id}_{safe_name}')
                proj_file.save(save_path)
                project_files.append(save_path)

        if not project_files:
            return jsonify({"error": "No se subieron archivos de proyecto válidos"}), 400

        font_family = request.form.get("fontFamily", "default")
        mode = request.form.get("mode", "unified")
        generate_ai = request.form.get("generateAI") == "true"
        export_format = request.form.get("exportFormat", "word")

        output_filename = f"Reporte_Mensual_Generado_V2_{task_id[:8]}.docx"
        output_path = os.path.join(OUTPUT_FOLDER, output_filename)
        
        thread = threading.Thread(target=worker_compilation_v2, args=(task_id, master_path, project_files, output_path, font_family, generate_ai, mode, export_format))
        thread.start()
        
        return jsonify({"taskId": task_id, "status": "Task Started V2"})

    except Exception as e:
        print(f"Error procesando peticion POST V2: {e}")
        return jsonify({"error": f"Error al procesar los archivos: {str(e)}"}), 500
    
    return jsonify({"taskId": task_id, "status": "Task Started V2"})

# ==========================================
# FUNCIONES V3 (COPIADO FORMATO MAESTRo)
# ==========================================

def worker_format_copy(task_id, master_path, project_files, output_folder):
    """
    Background job that runs Win32 COM locally to copy template formatting 
    to exactly N project files and returns a ZIP.
    """
    update_progress(task_id, 10, "Win32: Cargando Instancia V3...")
    print("Iniciando MS Word (V3) en servidor...")
    pythoncom.CoInitialize()
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.ScreenUpdating = False
    word.DisplayAlerts = False
    
    processed_files = []
    
    try:
        total_p = len(project_files)
        for i, doc_path in enumerate(project_files):
            update_progress(task_id, 10 + int(70 * (i/total_p)), f"Procesando archivo {i+1} de {total_p}...")
            
            # Nombre de salida dinámico
            base_name = os.path.basename(doc_path)
            name_without_ext = os.path.splitext(base_name)[0]
            out_filename = f"{name_without_ext}_Formateado.docx"
            out_path = os.path.join(output_folder, out_filename)
            
            try:
                # 1. Abrimos el documento maestro COMO NUEVO (Clonamos)
                new_doc = word.Documents.Add(Template=master_path)
                
                # Limpiar el contenido de texto del maestro (dejando encabezados, logos, formato)
                try:
                    new_doc.Content.Delete()
                except:
                    pass
                
                # 2. Abrimos el documento origen
                src_doc = word.Documents.Open(doc_path, ReadOnly=True, Visible=False)
                
                # 3. Limpiamos estilos básicos en su origen (sombreados) para evitar conflictos violentos
                for p in src_doc.Paragraphs:
                    if not p.Range.Information(12): # Si no está en tabla
                        p.Range.HighlightColorIndex = 0
                        p.Range.Font.Shading.BackgroundPatternColor = 16777215
                        p.Range.ParagraphFormat.Shading.BackgroundPatternColor = 16777215
                        
                # 4. Copiamos todo el contenido
                src_doc.Content.Copy()
                src_doc.Close(False)
                
                # 5. Pegamos en el nuevo doc (clonado del maestro) usando formato de destino
                selection = word.Selection
                selection.EndKey(Unit=6) # wdStory = 6
                selection.PasteAndFormat(19) # wdUseDestinationStylesRecovery
                
                # 6. AutoFit Tablas
                for table in new_doc.Tables:
                    try:
                        table.AutoFitBehavior(2)
                        table.Rows.Alignment = 1
                    except:
                        pass
                        
                # 7. Guardar y cerrar
                new_doc.SaveAs(out_path)
                new_doc.Close(False)
                processed_files.append(out_path)
                
            except Exception as e:
                print(f"Error V3 Copy/Paste en {doc_path}: {e}")
                update_progress(task_id, 0, f"Error en archivo {base_name}: {e}")
                try:
                    if 'src_doc' in locals(): src_doc.Close(False)
                    if 'new_doc' in locals(): new_doc.Close(False)
                except:
                    pass

        update_progress(task_id, 85, "Comprimiendo archivos generados en ZIP...")
        
        # Comprimir en carpeta contenedora
        zip_filename = f"Seguimientos_Formateados_V3_{uuid.uuid4().hex[:6]}.zip"
        zip_path = os.path.join(OUTPUT_FOLDER, zip_filename)
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file_path in processed_files:
                zipf.write(file_path, arcname=os.path.basename(file_path))
                
        # Limpieza de temporales opcional aquí

        update_progress(task_id, 100, "¡Proceso Completado!")
        
        global_tasks[task_id]['downloadUrls'] = {"zip": f"/download/{zip_filename}"}
        global_tasks[task_id]['success'] = True
        
    except Exception as e:
        global_tasks[task_id]['error'] = f"Excepción crítica durante worker v3: {e}"
        traceback.print_exc()
    finally:
        try:
            word.Quit()
        except:
            pass
        pythoncom.CoUninitialize()
        global_tasks[task_id]['completed'] = True

@app.route("/v3", methods=["GET"])
def index_v3():
    return render_template("index_v3.html")

@app.route("/api/generate_v3", methods=["POST"])
def generate_report_v3():
    print("Received async generation request V3")
    task_id = str(uuid.uuid4())
    global_tasks[task_id] = {
        "progress": 0, "status": "Iniciando colas en servidor V3...", "completed": False, 
        "success": False, "error": None, "error_ai": None, "downloadUrl": None, "aiDownloadUrl": None
    }

    try:
        if 'masterTemplate' not in request.files:
            return jsonify({"error": "No has subido la plantilla maestra"}), 400
        
        master_file = request.files['masterTemplate']
        master_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{task_id}_master_v3.docx')
        master_file.save(master_path)
        
        project_files = []
        # Usamos request.files porque subimos de indice dinamico
        for key in request.files:
            if key.startswith('projectFile_'):
                proj_file = request.files[key]
                if proj_file.filename:
                    safe_name = secure_filename(proj_file.filename)
                    save_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{task_id}_{safe_name}")
                    proj_file.save(save_path)
                    project_files.append(save_path)

        if not project_files:
            return jsonify({"error": "No se subieron archivos válidos"}), 400

        # Carpeta temporal para guardar los procesados antes de zipear
        output_folder_v3 = os.path.join(OUTPUT_FOLDER, f"v3_job_{task_id}")
        os.makedirs(output_folder_v3, exist_ok=True)
        
        thread = threading.Thread(target=worker_format_copy, args=(task_id, master_path, project_files, output_folder_v3))
        thread.start()
        
        return jsonify({"taskId": task_id, "status": "Task Started V3"})

    except Exception as e:
        print(f"Error procesando peticion POST V3: {e}")
        return jsonify({"error": f"Error al procesar archivos: {str(e)}"}), 500

# ==========================================
# FUNCIONES V4 (EXTRACCION DE PRODUCTOS LIMPIOS)
# ==========================================

def worker_extract_products(task_id, master_path, project_files, output_folder):
    """
    Background job that runs Win32 COM locally to extract specifically the 'Productos' 
    tables and content up to 'Reporte ejecución' from N project files.
    """
    update_progress(task_id, 10, "Win32: Cargando Instancia V4 (Limpieza Profunda)...")
    print("Iniciando MS Word (V4) en servidor...")
    pythoncom.CoInitialize()
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.ScreenUpdating = False
    word.DisplayAlerts = False
    
    processed_files = []
    
    try:
        total_p = len(project_files)
        for i, doc_path in enumerate(project_files):
            update_progress(task_id, 10 + int(70 * (i/total_p)), f"Extrayendo archivo {i+1} de {total_p}...")
            
            base_name = os.path.basename(doc_path)
            name_without_ext = os.path.splitext(base_name)[0]
            out_filename = f"{name_without_ext}_Limpio.docx"
            out_path = os.path.join(output_folder, out_filename)
            
            try:
                # 1. Abrimos origen y creamos destino
                src_doc = word.Documents.Open(doc_path, ReadOnly=True, Visible=False)
                new_doc = word.Documents.Add(Template=master_path)
                
                # Limpiar el contenido de texto del maestro (dejando encabezados, logos, formato)
                try:
                    new_doc.Content.Delete()
                except:
                    pass
                    
                dest_selection = word.Selection
                
                # Banderas de control
                found_productos = False
                
                # 2. Recorremos buscando la tabla de productos
                for table_idx in range(1, src_doc.Tables.Count + 1):
                    table = src_doc.Tables(table_idx)
                    try:
                        # Revisar solo celdas iniciales a ver si dice "Producto" o "Indicador de producto"
                        header_text = table.Cell(1,1).Range.Text.lower() + " " + table.Cell(1,2).Range.Text.lower()
                    except:
                        header_text = ""
                        
                    if "producto" in header_text or "metas" in header_text or "resultado" in header_text:
                        found_productos = True
                        start_range_pos = table.Range.Start
                        
                        # Buscar donde termina este bloque (Buscando "Reporte ejecución" o prox tabla)
                        end_range_pos = src_doc.Content.End
                        
                        search_range = src_doc.Range(Start=table.Range.End, End=src_doc.Content.End)
                        find_obj = search_range.Find
                        find_obj.Text = "Reporte ejecución"
                        find_obj.MatchCase = False
                        
                        if find_obj.Execute():
                            # Encontró el título tope
                            end_range_pos = search_range.Paragraphs.Last.Range.End
                        else:
                            # Buscar "Seguimiento a la ejecución"
                            find_obj.Text = "Seguimiento a la ejecución"
                            if find_obj.Execute():
                                end_range_pos = search_range.Paragraphs.Last.Range.End
                        
                        # Crear el rango definitivo a extraer
                        extract_range = src_doc.Range(Start=start_range_pos, End=end_range_pos)
                        
                        # Limpiar estilos antes de copiar
                        for p in extract_range.Paragraphs:
                            if not p.Range.Information(12): 
                                p.Range.HighlightColorIndex = 0
                                p.Range.Font.Shading.BackgroundPatternColor = 16777215
                                p.Range.ParagraphFormat.Shading.BackgroundPatternColor = 16777215
                        
                        # Copiar y pegar
                        extract_range.Copy()
                        
                        dest_selection.EndKey(Unit=6)
                        dest_selection.PasteAndFormat(19) # wdUseDestinationStylesRecovery
                        dest_selection.TypeParagraph()
                        
                src_doc.Close(False)
                
                # Formato final al destino
                for table in new_doc.Tables:
                    try:
                        table.AutoFitBehavior(2)
                        table.Rows.Alignment = 1
                    except:
                        pass
                        
                new_doc.SaveAs(out_path)
                new_doc.Close(False)
                
                # Si encontró algo, subimos. Si no encontró NA, también se sube vacío con el header del master
                processed_files.append(out_path)
                
            except Exception as e:
                print(f"Error V4 Extracción en {doc_path}: {e}")
                update_progress(task_id, 0, f"Error extracción en {base_name}: {e}")
                try:
                    if 'src_doc' in locals(): src_doc.Close(False)
                    if 'new_doc' in locals(): new_doc.Close(False)
                except:
                    pass

        update_progress(task_id, 85, "Comprimiendo archivos limpios en ZIP...")
        
        zip_filename = f"Seguimientos_Extraidos_V4_{uuid.uuid4().hex[:6]}.zip"
        zip_path = os.path.join(OUTPUT_FOLDER, zip_filename)
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file_path in processed_files:
                zipf.write(file_path, arcname=os.path.basename(file_path))
                
        update_progress(task_id, 100, "¡Extracción Limpia Completada!")
        
        global_tasks[task_id]['downloadUrls'] = {"zip": f"/download/{zip_filename}"}
        global_tasks[task_id]['success'] = True
        
    except Exception as e:
        global_tasks[task_id]['error'] = f"Excepción crítica durante worker v4: {e}"
        traceback.print_exc()
    finally:
        try:
            word.Quit()
        except:
            pass
        pythoncom.CoUninitialize()
        global_tasks[task_id]['completed'] = True

@app.route("/v4", methods=["GET"])
def index_v4():
    return render_template("index_v4.html")

@app.route("/api/generate_v4", methods=["POST"])
def generate_report_v4():
    print("Received async generation request V4")
    
    if 'masterTemplate' not in request.files:
        return jsonify({"error": "No has subido la plantilla maestra"}), 400
    
    master_file = request.files['masterTemplate']
    master_path = os.path.join(app.config['UPLOAD_FOLDER'], f'master_template_v4_{uuid.uuid4().hex[:5]}.docx')
    master_file.save(master_path)
    
    project_files = []
    for key in request.files:
        if key.startswith('projectFile_'):
            proj_file = request.files[key]
            if proj_file.filename:
                safe_name = secure_filename(proj_file.filename)
                save_path = os.path.join(app.config['UPLOAD_FOLDER'], f"v4_{uuid.uuid4().hex[:4]}_{safe_name}")
                proj_file.save(save_path)
                project_files.append(save_path)

    if not project_files:
        return jsonify({"error": "No se subieron archivos válidos"}), 400

    job_id = uuid.uuid4().hex[:8]
    job_output_folder = os.path.join(OUTPUT_FOLDER, f"v4_job_{job_id}")
    os.makedirs(job_output_folder, exist_ok=True)
    
    task_id = str(uuid.uuid4())
    global_tasks[task_id] = {
        "progress": 0, "status": "Iniciando motor de extracción...", "completed": False, 
        "success": False, "error": None, "downloadUrls": None
    }
    
    thread = threading.Thread(target=worker_extract_products, args=(task_id, master_path, project_files, job_output_folder))
    thread.start()
    
    return jsonify({"taskId": task_id, "status": "Task Started V4"})

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True, use_reloader=False)
